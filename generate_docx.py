from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_requirements_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('Documento de Requisitos Funcionales - SERVICED', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Este documento detalla los 56 requisitos funcionales del sistema SERVICED, estructurados por módulos operativos.').alignment = WD_ALIGN_PARAGRAPH.CENTER

    sections = {
        "1. Autenticación y Registro": [
            "El sistema permitirá el registro de nuevos usuarios con rol de 'Cliente'.",
            "El sistema permitirá el registro de nuevos usuarios con rol de 'Profesional'.",
            "El sistema validará que el correo electrónico ingresado tenga un formato válido.",
            "El sistema verificará que el correo electrónico no esté registrado previamente.",
            "El sistema requerirá una contraseña de al menos 8 caracteres.",
            "El sistema validará la presencia de al menos una letra mayúscula en la contraseña.",
            "El sistema validará la presencia de al menos un carácter especial en la contraseña.",
            "El sistema permitirá el inicio de sesión seguro mediante credenciales (email/password)."
        ],
        "2. Gestión de Perfil y Cuenta": [
            "El usuario podrá actualizar su nombre completo desde el perfil.",
            "El usuario podrá actualizar su número de teléfono.",
            "El usuario podrá cambiar su contraseña actual desde la configuración.",
            "El usuario podrá subir o cambiar su foto de perfil.",
            "El sistema permitirá la recuperación de contraseña vía correo electrónico.",
            "El usuario podrá visualizar el rol asignado a su cuenta.",
            "El sistema permitirá al usuario cerrar sesión de forma segura.",
            "El usuario podrá configurar su ubicación para mejorar la relevancia de servicios."
        ],
        "3. Búsqueda y Exploración de Servicios": [
            "El sistema permitirá buscar servicios por palabras clave.",
            "El sistema permitirá filtrar servicios por categoría.",
            "El sistema permitirá ordenar los resultados por relevancia o calificación.",
            "El cliente podrá ver el perfil detallado de un profesional antes de contratar.",
            "El sistema mostrará la calificación promedio de cada servicio.",
            "El cliente podrá ver reseñas de otros usuarios en el detalle del servicio.",
            "El sistema mostrará el precio base de cada servicio publicado.",
            "El cliente podrá visualizar la disponibilidad básica del profesional.",
            "El sistema permitirá guardar servicios en una sección de 'Favoritos'.",
            "El sistema mostrará servicios sugeridos basados en búsquedas previas."
        ],
        "4. Gestión de Solicitudes y Pedidos": [
            "El cliente podrá realizar una solicitud de servicio a un profesional específico.",
            "El profesional podrá aceptar una solicitud de servicio entrante.",
            "El profesional podrá rechazar una solicitud de servicio con una breve justificación.",
            "El sistema permitirá al cliente cancelar una solicitud pendiente.",
            "El profesional podrá marcar un servicio como 'En Progreso'.",
            "El profesional podrá marcar un servicio como 'Completado'.",
            "El sistema registrará la fecha y hora de creación de cada solicitud.",
            "El cliente podrá ver el historial de todos sus pedidos realizados.",
            "El profesional podrá ver el historial de trabajos finalizados.",
            "El sistema enviará una confirmación al cliente cuando un servicio sea aceptado."
        ],
        "5. Comunicación y Chat": [
            "El sistema proporcionará un chat en tiempo real entre cliente y profesional.",
            "El sistema guardará el historial de mensajes de cada conversación.",
            "Los usuarios podrán ver el estado de conexión del otro usuario en el chat.",
            "El sistema permitirá enviar mensajes de texto y caracteres básicos.",
            "El usuario recibirá una notificación visual al recibir un nuevo mensaje.",
            "El sistema permitirá ver la fecha y hora de los mensajes enviados."
        ],
        "6. Panel Administrativo": [
            "El administrador podrá visualizar todos los usuarios registrados en el sistema.",
            "El administrador tendrá la capacidad de desactivar cuentas de usuario.",
            "El administrador podrá eliminar usuarios de forma permanente.",
            "El administrador podrá crear nuevas categorías de servicio.",
            "El administrador podrá editar o eliminar categorías existentes.",
            "El administrador podrá visualizar reportes de actividad global.",
            "El administrador podrá supervisar las solicitudes de servicios activas.",
            "El administrador podrá gestionar la configuración general del sitio."
        ],
        "7. Dashboard y Notificaciones": [
            "El profesional podrá visualizar un dashboard con el resumen de sus ganancias.",
            "El profesional podrá ver el número total de servicios realizados en el mes.",
            "El sistema emitirá notificaciones sonoras al recibir una nueva solicitud.",
            "El cliente recibirá una notificación al ser calificado por un profesional.",
            "El sistema mostrará un contador de notificaciones no leídas en el sidebar.",
            "El sistema permitirá alternar entre modo claro y modo oscuro en el panel."
        ]
    }

    req_number = 1
    for section_title, requirements in sections.items():
        doc.add_heading(section_title, level=1)
        for req in requirements:
            p = doc.add_paragraph(style='List Number')
            run = p.add_run(f"RF-{req_number:02d}: {req}")
            req_number += 1

    file_path = r'c:\Users\alvar\Downloads\serviced2.1\serviced2.2\docs\requisitos_funcionales.docx'
    doc.save(file_path)
    print(f"Documento generado exitosamente en: {file_path}")

if __name__ == "__main__":
    create_requirements_doc()
