from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_presentation_order(path):
    doc = Document()
    
    # Título Principal
    title = doc.add_heading('Estructura de Presentación Final: SERVICED (ADSO)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Este documento define el orden lógico y los puntos clave para la sustentación final del proyecto, asegurando el cumplimiento del checklist de evaluación.')

    # Bloque 1: Introducción y Propuesta
    doc.add_heading('Fase 1: El Pitch y la Propuesta', level=1)
    items1 = [
        ('1. Presentación del Pitch (1 min)', 'Apertura de alto impacto con los nombres del equipo y propuesta de valor.'),
        ('9. Propuestas Técnicas de Servicios TI', 'Presentación de costos de desarrollo, inversión inicial y presupuesto operativo mensual (COP).'),
        ('8. Prototipo Gráfico', 'Muestra del diseño en Figma/Mockups para dar contexto visual antes del código.')
    ]
    for item, desc in items1:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.bold = True
        p.add_run(f': {desc}')

    # Bloque 2: Análisis y Diseño
    doc.add_heading('Fase 2: Documentación Técnica y Requisitos', level=1)
    items2 = [
        ('2. Informe de Especificación de Requisitos', 'Validación de RFs y RNFs del sistema.'),
        ('3. Informes de Análisis y Diseño (Artefactos)', 'Diagramas UML, casos de uso, diagramas de clases y arquitectura.'),
        ('13. Manuales Técnicos', 'Entrega de Manual de Usuario, Técnico, Despliegue y Documentación de la API.')
    ]
    for item, desc in items2:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.bold = True
        p.add_run(f': {desc}')

    # Bloque 3: Base de Datos y Backend
    doc.add_heading('Fase 3: Implementación del Lado del Servidor', level=1)
    items3 = [
        ('4. SQL de la Base de Datos', 'Explicación del modelo E-R y scripts de creación de tablas.'),
        ('6. Código del Software (Backend)', 'Demostración de FastAPI, modelos, esquemas y lógica de negocio integrada.'),
        ('12. Herramientas de Terceros', 'Demostración de reportes PDF/Excel, envío de correos o integraciones de IA.')
    ]
    for item, desc in items3:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.bold = True
        p.add_run(f': {desc}')

    # Bloque 4: Frontend y Calidad
    doc.add_heading('Fase 4: Interfaz y Calidad de Código', level=1)
    items4 = [
        ('5. Código del Software (Frontend)', 'Interfaz funcional y responsiva integrada con el backend.'),
        ('7. Explicación de Calidad de Código', 'Muestra de indentación, comentarios, uso de linters y estándares ADSO.'),
        ('15. Desempeño Individual', 'Manejo del código fuente por parte de cada integrante durante la explicación.')
    ]
    for item, desc in items4:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.bold = True
        p.add_run(f': {desc}')

    # Bloque 5: Pruebas y Despliegue
    doc.add_heading('Fase 5: Validación y Puesta en Marcha', level=1)
    items5 = [
        ('10. Informe de Resultados de Pruebas', 'Evidencia de pruebas unitarias, de integración y funcionales.'),
        ('11. Funcionamiento y Verificación', 'Muestra de validaciones, manejo de errores (try/catch), alertas y seguridad.'),
        ('16. Despliegue de la Aplicación', 'App funcionando en entorno de producción (Backend y Frontend).'),
        ('14. Dominio del Tema', 'Evaluación continua del manejo del tiempo y conocimiento técnico.')
    ]
    for item, desc in items5:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.bold = True
        p.add_run(f': {desc}')

    # Bloque 6: Cierre Administrativo
    doc.add_heading('Fase 6: Cierre y Entrega Legal', level=1)
    items6 = [
        ('17. Enlace Público y Repositorio', 'Entrega de enlaces de Github, releases en ZIP y correos de coordinación.')
    ]
    for item, desc in items6:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.bold = True
        p.add_run(f': {desc}')

    doc.save(path)

if __name__ == "__main__":
    path = r"c:/Users/alvar/Downloads/serviced2.1/serviced2.2/docs/orden_presentacion.docx"
    create_presentation_order(path)
    print(f"Archivo generado en: {path}")
