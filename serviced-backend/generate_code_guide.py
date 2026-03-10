from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_code_guide(path):
    doc = Document()
    
    # Título
    title = doc.add_heading('Guía de Estudio y Navegación del Código: SERVICED', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Este documento te ayudará a entender la ubicación exacta de cada componente de tu proyecto para que puedas practicar y dominar la arquitectura ADSO.')

    # Sección 1: El Backend (FastAPI)
    doc.add_heading('1. El Backend: La Lógica y los Datos', level=1)
    
    doc.add_heading('¿Dónde está la Conexión a la Base de Datos?', level=2)
    doc.add_paragraph('La configuración de la conexión (SQLite/PostgreSQL) y la sesión de la base de datos están en:')
    doc.add_paragraph('serviced-backend/app/db/session.py', style='List Bullet')
    
    doc.add_heading('¿Dónde están los Modelos de la DB?', level=2)
    doc.add_paragraph('Las clases que definen las tablas de la base de datos están en:')
    doc.add_paragraph('serviced-backend/app/models/', style='List Bullet')
    doc.add_paragraph('Ejemplo: user.py, service.py, request.py', style='List Bullet')

    doc.add_heading('¿Dónde está el CRUD (Create, Read, Update, Delete)?', level=2)
    doc.add_paragraph('La lógica que interactúa directamente con la base de datos se encuentra en los repositorios:')
    doc.add_paragraph('serviced-backend/app/repositories/', style='List Bullet')
    doc.add_paragraph('Aquí es donde verás las funciones como "db.add()", "db.commit()", "db.query()".', style='Body Text')

    doc.add_heading('¿Dónde están los Endpoints (Rutas)?', level=2)
    doc.add_paragraph('Aquí llegan las peticiones del frontend:')
    doc.add_paragraph('serviced-backend/app/api/v1/endpoints/', style='List Bullet')

    # Sección 2: El Frontend (HTML/JS)
    doc.add_heading('2. El Frontend: Interfaz y Llamadas (Fetch)', level=1)
    doc.add_paragraph('Tienes tres carpetas principales para el frontend:')
    doc.add_paragraph('serviced-users: Interfaz para clientes.', style='List Bullet')
    doc.add_paragraph('serviced-provider: Interfaz para profesionales.', style='List Bullet')
    doc.add_paragraph('serviced-admin: Panel de administración.', style='List Bullet')

    doc.add_heading('¿Dónde están los Fetch (Peticiones al Backend)?', level=2)
    doc.add_paragraph('Busca los archivos .js o los bloques <script> dentro de los HTML de cada carpeta.')
    doc.add_paragraph('Ejemplo: serviced-users/user-chat.js contiene los fetch para el sistema de chat.', style='List Bullet')
    doc.add_paragraph('Busca la palabra clave "fetch(" para ver a qué URL del backend se está llamando.', style='Body Text')

    # Sección 3: Flujo de Práctica Sugerido
    doc.add_heading('3. Cómo practicar: Sigue el Hilo', level=1)
    doc.add_paragraph('Para dominar tu código, elige una funcionalidad (ej. Crear una Solicitud) y sigue este camino:')
    
    steps = [
        ('FRONTEND (JS)', 'Busca el "fetch" en el archivo JS que envía los datos.'),
        ('BACKEND (API)', 'Ubica la ruta en "app/api/v1/endpoints" que recibe ese fetch.'),
        ('BACKEND (CRUD)', 'Mira qué función del repositorio se llama para guardar en la DB.'),
        ('DATABASE (MODEL)', 'Verifica cómo está definida esa tabla en "app/models".')
    ]
    
    for i, (loc, desc) in enumerate(steps, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'Paso {i}: {loc}')
        run.bold = True
        p.add_run(f' - {desc}')

    # Sección 4: Recomendaciones Finales
    doc.add_heading('4. Tips para la Sustentación', level=1)
    doc.add_paragraph('Manejo de errores: Busca los bloques "try...except" en el backend y ".catch()" en el frontend.', style='List Bullet')
    doc.add_paragraph('Validaciones (Pydantic): Revisa la carpeta "app/schemas" para ver cómo se validan los datos que llegan de la API.', style='List Bullet')

    doc.save(path)

if __name__ == "__main__":
    path = r"c:/Users/alvar/Downloads/serviced2.1/serviced2.2/docs/guia_estudio_codigo.docx"
    create_code_guide(path)
    print(f"Guía de código generada en: {path}")
