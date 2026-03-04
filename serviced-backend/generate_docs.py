from docx import Document
from docx.shared import Pt
import os

def create_pitch_docx(path):
    doc = Document()
    doc.add_heading('Guion del Pitch: SERVICED (1 Minuto)', 0)
    
    p = doc.add_paragraph()
    run = p.add_run('Integrantes: Nicolás Álvarez Macías, Angel Favian Romero, Sebastian Benitez.')
    run.bold = True
    
    doc.add_heading('0:00 - 0:10 Presentación', level=2)
    doc.add_paragraph('"Hola a todos. Somos Nicolás Álvarez, Angel Romero y Sebastian Benitez, y hoy les presentamos SERVICED, la solución definitiva para conectar el talento con la necesidad."')
    
    doc.add_heading('0:10 - 0:20 ¿Qué es?', level=2)
    doc.add_paragraph('"SERVICED es un marketplace de servicios profesionales diseñado para que clientes encuentren expertos de confianza en segundos, y profesionales escalen sus negocios de forma digital."')
    
    doc.add_heading('0:20 - 0:40 Logros y Objetivos', level=2)
    doc.add_paragraph('"Se han logrado hitos clave: implementamos dashboards robustos para Clientes, Proveedores y Administradores. Además, contamos con un sistema de chat en tiempo real y notificaciones instantáneas que garantizan una comunicación fluida y segura."')
    
    doc.add_heading('0:40 - 0:50 Arquitectura Tecnológica', level=2)
    doc.add_paragraph('"Nuestra infraestructura se basa en una arquitectura escalable: un backend potente con FastAPI y una base de datos SQLite optimizada, garantizando velocidad y alta disponibilidad en una interfaz totalmente responsiva."')
    
    doc.add_heading('0:50 - 1:00 Costos y Cierre', level=2)
    doc.add_paragraph('"Nuestra estructura financiera es clara: una inversión inicial de $40 millones de pesos para desarrollo y un costo operativo mensual de apenas $500,000 pesos. Estamos listos para revolucionar el mercado de servicios. ¡Esto es SERVICED!"')
    
    doc.save(path)

def create_costs_docx(path):
    doc = Document()
    doc.add_heading('Análisis de Costos Detallado: SERVICED', 0)
    
    doc.add_heading('1. Costos de Desarrollo (Inversión Inicial)', level=1)
    table1 = doc.add_table(rows=1, cols=3)
    table1.style = 'Table Grid'
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = 'Categoría'
    hdr_cells[1].text = 'Descripción'
    hdr_cells[2].text = 'Costo (COP)'
    
    data1 = [
        ('Personal', 'Equipo de Desarrollo (3 meses)', '$36.000.000'),
        ('Diseño', 'Branding y UI/UX', '$2.500.000'),
        ('Legal', 'Constitución y Términos', '$1.800.000'),
        ('TOTAL', '', '$40.300.000')
    ]
    
    for cat, desc, cost in data1:
        row_cells = table1.add_row().cells
        row_cells[0].text = cat
        row_cells[1].text = desc
        row_cells[2].text = cost

    doc.add_paragraph()
    doc.add_heading('2. Costos Operativos Mensuales (Mantenimiento)', level=1)
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Table Grid'
    hdr_cells = table2.rows[0].cells
    hdr_cells[0].text = 'Concepto'
    hdr_cells[1].text = 'Detalle'
    hdr_cells[2].text = 'Costo (COP)'
    
    data2 = [
        ('Infraestructura', 'Servidor Cloud', '$80.000'),
        ('Base de Datos', 'DB Gestionada', '$60.000'),
        ('Diseño', 'Figma Professional', '$65.000'),
        ('Herramientas', 'GitHub Copilot / IDEs', '$40.000'),
        ('Dominio/SSL', 'Seguridad y Web', '$10.000'),
        ('Mensajería', 'APIs de Notificación', '$40.000'),
        ('Marketing', 'Ads / RRSS', '$200.000'),
        ('TOTAL MENSUAL', '', '$495.000')
    ]
    
    for conc, det, cost in data2:
        row_cells = table2.add_row().cells
        row_cells[0].text = conc
        row_cells[1].text = det
        row_cells[2].text = cost

    doc.add_paragraph()
    doc.add_heading('3. Escenario Mensual con Personal', level=1)
    doc.add_paragraph('Total Gastos Mensuales (Nómina + Operación): $12.495.000 COP')
    
    doc.save(path)

if __name__ == "__main__":
    docs_dir = r"c:/Users/alvar/Downloads/serviced2.1/serviced2.2/docs"
    if not os.path.exists(docs_dir):
        os.makedirs(docs_dir)
    
    create_pitch_docx(os.path.join(docs_dir, "pitch_guion.docx"))
    create_costs_docx(os.path.join(docs_dir, "tabla_costos.docx"))
    print("Documentos generados exitosamente.")
