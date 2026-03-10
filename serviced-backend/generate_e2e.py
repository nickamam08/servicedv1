from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_e2e_guide(path):
    doc = Document()
    
    # Título
    title = doc.add_heading('Guía Paso a Paso: Pruebas End-to-End (E2E) - SERVICED', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Las pruebas End-to-End (E2E) validan que el sistema completo funcione correctamente desde la interfaz de usuario hasta la base de datos, simulando el comportamiento real de un usuario.')

    # Sección 1: Introducción
    doc.add_heading('1. ¿Qué es una Prueba End-to-End?', level=1)
    doc.add_paragraph('Es una metodología de prueba que verifica el flujo funcional de una aplicación de principio a fin. El objetivo es asegurar que todos los componentes (Frontend, Backend, Base de Datos, APIs) se comuniquen y operen correctamente bajo escenarios reales.')

    # Sección 2: Preparación del Entorno
    doc.add_heading('2. Preparación del Entorno', level=1)
    p = doc.add_paragraph('Antes de iniciar, asegúrese de que:')
    doc.add_paragraph('El backend esté en ejecución (FastAPI / Uvicorn).', style='List Bullet')
    doc.add_paragraph('El frontend esté accesible (Vite / Live Server).', style='List Bullet')
    doc.add_paragraph('La base de datos (SQLite) tenga datos iniciales o esté limpia para pruebas.', style='List Bullet')

    # Sección 3: El Flujo E2E Maestro (Caso de Uso: Registro y Solicitud)
    doc.add_heading('3. Flujo Crítico a Probar', level=1)
    doc.add_paragraph('Para SERVICED, el flujo End-to-End más importante es el siguiente:')
    
    steps = [
        ('Registro de Cliente', 'El usuario ingresa datos en el formulario de registro del frontend.'),
        ('Autenticación', 'Inicia sesión y recibe un token JWT desde el backend.'),
        ('Búsqueda de Servicio', 'Navega por las categorías y selecciona un profesional.'),
        ('Creación de Solicitud', 'Llena el formulario de solicitud y la envía.'),
        ('Notificación al Proveedor', 'El backend procesa la solicitud y genera una notificación en tiempo real.'),
        ('Gestión del Proveedor', 'El profesional acepta la solicitud desde su propio dashboard.'),
        ('Chat en Vivo', 'Cliente y profesional intercambian mensajes para coordinar.')
    ]
    
    for i, (step, desc) in enumerate(steps, 1):
        doc.add_heading(f'Paso {i}: {step}', level=2)
        doc.add_paragraph(desc)

    # Sección 4: Herramientas Sugeridas
    doc.add_heading('4. Herramientas para Automatizar E2E', level=1)
    doc.add_paragraph('Si desea automatizar estas pruebas en el futuro, se recomiendan:')
    doc.add_paragraph('Cypress: La herramienta más popular para pruebas web rápidas.', style='List Bullet')
    doc.add_paragraph('Playwright: Desarrollada por Microsoft, ideal para pruebas multiplataforma.', style='List Bullet')
    doc.add_paragraph('Selenium: El estándar de la industria para navegadores antiguos.', style='List Bullet')

    # Sección 5: Checklist de Verificación
    doc.add_heading('5. Checklist de Éxito', level=1)
    checks = [
        '¿Se guardó el registro correctamente en la base de datos SQL?',
        '¿El frontend mostró el mensaje de éxito o alerta adecuado?',
        '¿Las peticiones a la API devolvieron código 200 OK?',
        '¿El estado de la solicitud cambió correctamente en el dashboard?'
    ]
    for check in checks:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'[ ] {check}')

    doc.save(path)

if __name__ == "__main__":
    path = r"c:/Users/alvar/Downloads/serviced2.1/serviced2.2/docs/guia_e2e_paso_a_paso.docx"
    create_e2e_guide(path)
    print(f"Guía E2E generada en: {path}")
